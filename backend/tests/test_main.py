"""
Tests for the AI Career Copilot hackathon demo backend.

Every call to Mesh API is mocked (via unittest.mock.patch) — these tests
never make real network calls, never cost real money, and never require a
real API key. This lets you run the full suite offline, in CI, or right
before a demo without worrying about spend.

Since /api/parse-resume and /api/skill-gap are now protected routes that
also write to the database, this suite uses FastAPI's dependency_overrides
to substitute a fake logged-in user and a fake (mocked) database session —
the standard way to test protected/DB-backed routes without needing a real
JWT token or a real database connection.

Run:
    pip install pytest httpx
    pytest
"""

import io
import json
import uuid
from datetime import datetime, timezone
from types import SimpleNamespace
from unittest.mock import patch, MagicMock

import pytest
from fastapi.testclient import TestClient

import main
import auth
from database import get_db

client = TestClient(main.app)


def make_fake_response(content_str: str, prompt_tokens: int = 10, completion_tokens: int = 20):
    """Builds an object shaped like Mesh/OpenAI's chat completion response,
    just enough for main.call_model() to read .choices[0].message.content
    and .usage.prompt_tokens / .usage.completion_tokens."""
    return SimpleNamespace(
        choices=[SimpleNamespace(message=SimpleNamespace(content=content_str))],
        usage=SimpleNamespace(prompt_tokens=prompt_tokens, completion_tokens=completion_tokens),
    )


# --- Fake logged-in user + fake DB session, applied to every test ---
# This is the standard FastAPI pattern for testing protected/DB-backed
# routes: override the dependency functions themselves rather than trying
# to fake a real JWT or spin up a real database.

FAKE_USER_ID = uuid.uuid4()

fake_user = SimpleNamespace(
    id=FAKE_USER_ID,
    name="Test User",
    email="test@example.com",
    auth_provider="password",
    google_sub=None,
    created_at=datetime.now(timezone.utc),
)


def override_get_current_user():
    return fake_user


@pytest.fixture(autouse=True)
def override_dependencies():
    """Applied automatically to every test in this file. Overrides auth
    so every request is treated as fake_user being logged in, and
    overrides the database session with a MagicMock so no real database
    connection is ever needed. Cleans up after each test."""
    mock_db = MagicMock()
    # By default, any .query(...).filter(...).first() call returns None
    # (as if nothing was found) — individual tests can override this
    # return value on mock_db if they need to simulate a found record.
    mock_db.query.return_value.filter.return_value.first.return_value = None

    def override_get_db():
        yield mock_db

    main.app.dependency_overrides[auth.get_current_user] = override_get_current_user
    main.app.dependency_overrides[get_db] = override_get_db

    yield mock_db  # tests can access/customize the mock via this fixture

    main.app.dependency_overrides.clear()


# --- Health check (not a protected route, no override needed) ---

def test_health_check():
    resp = client.get("/")
    assert resp.status_code == 200
    assert resp.json()["status"] == "ok"


# --- Resume parsing: happy path ---

def test_parse_resume_txt_success():
    fake_json = json.dumps({
        "name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "",
        "education": [],
        "skills": ["Python", "SQL"],
        "experience": [],
        "projects": [],
        "certifications": [],
    })
    fake_response = make_fake_response(fake_json)

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/parse-resume",
            files={"file": ("resume.txt", b"Jane Doe\nSkills: Python, SQL", "text/plain")},
        )

    assert resp.status_code == 200
    body = resp.json()
    assert body["data"]["name"] == "Jane Doe"
    assert "Python" in body["data"]["skills"]
    assert body["usage"]["input_tokens"] == 10
    assert body["usage"]["output_tokens"] == 20
    # Confirms the endpoint actually attempted to persist the resume —
    # add/commit/refresh should have been called on our mock db session.
    assert "resume_id" in body


def test_parse_resume_docx_success(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Skills: Python, SQL")
    docx_path = tmp_path / "resume.docx"
    doc.save(docx_path)
    docx_bytes = docx_path.read_bytes()

    fake_json = json.dumps({
        "name": "Jane Doe", "email": "", "phone": "",
        "education": [], "skills": ["Python", "SQL"],
        "experience": [], "projects": [], "certifications": [],
    })
    fake_response = make_fake_response(fake_json)

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/parse-resume",
            files={"file": (
                "resume.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )},
        )

    assert resp.status_code == 200
    assert resp.json()["data"]["name"] == "Jane Doe"


# --- Resume parsing: validation / error handling ---

def test_parse_resume_rejects_unsupported_file_type():
    resp = client.post(
        "/api/parse-resume",
        files={"file": ("resume.xyz", b"some content", "application/octet-stream")},
    )
    assert resp.status_code == 400
    assert "Unsupported file type" in resp.json()["detail"]


def test_parse_resume_rejects_empty_file():
    resp = client.post(
        "/api/parse-resume",
        files={"file": ("resume.txt", b"", "text/plain")},
    )
    assert resp.status_code == 400


def test_parse_resume_rejects_oversized_file():
    big_content = b"a" * (main.MAX_FILE_SIZE_BYTES + 1)
    resp = client.post(
        "/api/parse-resume",
        files={"file": ("resume.txt", big_content, "text/plain")},
    )
    assert resp.status_code == 413


def test_parse_resume_handles_invalid_json_from_model():
    # Simulates the model returning something that isn't valid JSON —
    # the endpoint should fail cleanly (502) instead of crashing or
    # passing garbage back to the frontend.
    fake_response = make_fake_response("Sorry, I can't help with that.")

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/parse-resume",
            files={"file": ("resume.txt", b"Some resume text", "text/plain")},
        )

    assert resp.status_code == 502


# --- Resume parsing: authentication ---

def test_parse_resume_requires_auth():
    # Remove the auth override just for this one test, to confirm the
    # route is genuinely protected when no valid user is provided.
    main.app.dependency_overrides.pop(auth.get_current_user, None)
    try:
        resp = client.post(
            "/api/parse-resume",
            files={"file": ("resume.txt", b"Some resume text", "text/plain")},
        )
        assert resp.status_code == 401
    finally:
        # Restore the override so it doesn't leak into other tests
        main.app.dependency_overrides[auth.get_current_user] = override_get_current_user


# --- Skill gap: happy path ---

def test_skill_gap_success():
    fake_json = json.dumps({
        "missing_skills": [
            {"skill": "DSA", "priority": "high", "why_it_matters": "Interviews"}
        ],
        "roadmap": [
            {"step": 1, "title": "Learn DSA", "description": "...", "estimated_weeks": 4, "resources": ["LeetCode"]}
        ],
        "recommended_projects": [
            {"title": "Project X", "description": "...", "skills_practiced": ["Python"]}
        ],
        "job_readiness_score": 42,
    })
    fake_response = make_fake_response(fake_json)

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/skill-gap",
            data={"role": "SDE", "skills": "Python, SQL"},
        )

    assert resp.status_code == 200
    body = resp.json()
    assert body["data"]["job_readiness_score"] == 42
    assert body["data"]["missing_skills"][0]["skill"] == "DSA"


def test_skill_gap_links_to_resume_when_owned(override_dependencies):
    # Simulate a resume that DOES belong to the logged-in fake_user, so
    # the endpoint's ownership check passes and it attempts to save a
    # linked ResumeAnalysis row.
    mock_db = override_dependencies
    fake_resume = SimpleNamespace(id=uuid.uuid4(), user_id=FAKE_USER_ID)
    mock_db.query.return_value.filter.return_value.first.return_value = fake_resume

    fake_json = json.dumps({
        "missing_skills": [], "roadmap": [], "recommended_projects": [],
        "job_readiness_score": 70,
    })
    fake_response = make_fake_response(fake_json)

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/skill-gap",
            data={"role": "SDE", "skills": "Python", "resume_id": str(fake_resume.id)},
        )

    assert resp.status_code == 200
    # add() should have been called to persist the ResumeAnalysis record
    assert mock_db.add.called


def test_skill_gap_malformed_resume_id_still_succeeds():
    # Regression test: a malformed (non-UUID) resume_id used to cause a
    # raw PostgreSQL error (500) because it was compared directly against
    # a UUID column. Since resume_id is optional metadata here (not the
    # endpoint's primary resource), the correct behavior is to silently
    # skip linking and still return the skill-gap result successfully.
    fake_json = json.dumps({
        "missing_skills": [], "roadmap": [], "recommended_projects": [],
        "job_readiness_score": 55,
    })
    fake_response = make_fake_response(fake_json)

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/skill-gap",
            data={"role": "SDE", "skills": "Python", "resume_id": "not-a-real-uuid"},
        )

    assert resp.status_code == 200


# --- Skill gap: validation ---

def test_skill_gap_rejects_empty_role():
    resp = client.post("/api/skill-gap", data={"role": "   ", "skills": "Python"})
    assert resp.status_code == 400


def test_skill_gap_rejects_oversized_role():
    resp = client.post(
        "/api/skill-gap",
        data={"role": "x" * (main.MAX_ROLE_LENGTH + 1), "skills": "Python"},
    )
    assert resp.status_code == 400


def test_skill_gap_rejects_oversized_skills():
    resp = client.post(
        "/api/skill-gap",
        data={"role": "SDE", "skills": "x" * (main.MAX_SKILLS_LENGTH + 1)},
    )
    assert resp.status_code == 400


def test_skill_gap_requires_auth():
    main.app.dependency_overrides.pop(auth.get_current_user, None)
    try:
        resp = client.post("/api/skill-gap", data={"role": "SDE", "skills": "Python"})
        assert resp.status_code == 401
    finally:
        main.app.dependency_overrides[auth.get_current_user] = override_get_current_user


# --- AI Interview Coach ---

def test_interview_start_success(override_dependencies):
    mock_db = override_dependencies
    fake_response = make_fake_response("Tell me about a challenging project you worked on.")

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/interview/start",
            data={"target_role": "SDE - Entry Level"},
        )

    assert resp.status_code == 200
    body = resp.json()
    assert body["question_number"] == 1
    assert body["total_questions"] == 4
    assert body["is_final"] is False
    assert "message" in body
    assert mock_db.add.called
    assert mock_db.commit.called


def test_interview_start_rejects_empty_role():
    resp = client.post("/api/interview/start", data={"target_role": "   "})
    assert resp.status_code == 400


def test_interview_start_requires_auth():
    main.app.dependency_overrides.pop(auth.get_current_user, None)
    try:
        resp = client.post("/api/interview/start", data={"target_role": "SDE"})
        assert resp.status_code == 401
    finally:
        main.app.dependency_overrides[auth.get_current_user] = override_get_current_user


def test_interview_answer_mid_session(override_dependencies):
    mock_db = override_dependencies
    fake_interview = SimpleNamespace(
        id=uuid.uuid4(),
        user_id=FAKE_USER_ID,
        target_role="SDE",
        transcript=[
            {"role": "system", "content": "You are an interviewer."},
            {"role": "user", "content": "I'm ready."},
            {"role": "assistant", "content": "Tell me about a challenge you faced."},
        ],
        question_number=1,
        status="in_progress",
    )
    mock_db.query.return_value.filter.return_value.first.return_value = fake_interview

    fake_response = make_fake_response("Good answer! Next: explain how hash maps work.")

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/interview/answer",
            data={"interview_id": str(fake_interview.id), "answer": "I built a Flask app with a tricky schema."},
        )

    assert resp.status_code == 200
    body = resp.json()
    assert body["is_final"] is False
    assert body["question_number"] == 2
    assert mock_db.commit.called


def test_interview_answer_final_question_produces_report(override_dependencies):
    mock_db = override_dependencies
    fake_interview = SimpleNamespace(
        id=uuid.uuid4(),
        user_id=FAKE_USER_ID,
        target_role="SDE",
        transcript=[
            {"role": "system", "content": "You are an interviewer."},
            {"role": "user", "content": "I'm ready."},
        ],
        question_number=4,  # already on the last question
        status="in_progress",
    )
    mock_db.query.return_value.filter.return_value.first.return_value = fake_interview

    fake_summary_json = json.dumps({
        "final_feedback": "Solid closing answer.",
        "score": 8,
        "strengths": ["Clear communication", "Good problem-solving approach"],
        "improvement_points": ["Practice more DSA", "Be more concise"],
    })
    fake_response = make_fake_response(fake_summary_json)

    with patch.object(main.client.chat.completions, "create", return_value=fake_response):
        resp = client.post(
            "/api/interview/answer",
            data={"interview_id": str(fake_interview.id), "answer": "My final answer to the last question."},
        )

    assert resp.status_code == 200
    body = resp.json()
    assert body["is_final"] is True
    assert body["score"] == 8
    assert "Clear communication" in body["strengths"]
    assert fake_interview.status == "completed"
    assert mock_db.add.called
    assert mock_db.commit.called


def test_interview_answer_rejects_empty_answer():
    resp = client.post(
        "/api/interview/answer",
        data={"interview_id": str(uuid.uuid4()), "answer": "   "},
    )
    assert resp.status_code == 400


def test_interview_answer_not_found(override_dependencies):
    mock_db = override_dependencies
    mock_db.query.return_value.filter.return_value.first.return_value = None

    resp = client.post(
        "/api/interview/answer",
        data={"interview_id": str(uuid.uuid4()), "answer": "Some answer"},
    )
    assert resp.status_code == 404


def test_interview_answer_requires_auth():
    main.app.dependency_overrides.pop(auth.get_current_user, None)
    try:
        resp = client.post(
            "/api/interview/answer",
            data={"interview_id": str(uuid.uuid4()), "answer": "test"},
        )
        assert resp.status_code == 401
    finally:
        main.app.dependency_overrides[auth.get_current_user] = override_get_current_user


def test_interview_report_not_completed(override_dependencies):
    mock_db = override_dependencies
    fake_interview = SimpleNamespace(
        id=uuid.uuid4(),
        user_id=FAKE_USER_ID,
        status="in_progress",
        report=None,
    )
    mock_db.query.return_value.filter.return_value.first.return_value = fake_interview

    resp = client.get(f"/api/interview/{fake_interview.id}/report")
    assert resp.status_code == 400


def test_interview_report_malformed_id_returns_404_not_500():
    # Regression test: a malformed (non-UUID-shaped) interview_id used to
    # cause a raw PostgreSQL error (surfacing as 500) because the ID was
    # compared directly against a UUID column before being validated.
    # Real database access isn't even needed to prove this - the ID should
    # be rejected as invalid before any query happens.
    resp = client.get("/api/interview/not-a-real-uuid-at-all/report")
    assert resp.status_code == 404


def test_interview_answer_malformed_id_returns_404_not_500():
    resp = client.post(
        "/api/interview/answer",
        data={"interview_id": "also-not-a-uuid", "answer": "test answer"},
    )
    assert resp.status_code == 404


# --- Security-relevant behavior ---

def test_parse_resume_does_not_leak_raw_exception_details():
    # A corrupted "PDF" (just random bytes with a .pdf name) should trigger
    # our sanitized error message, not raise a raw internal traceback/string
    # back to the client.
    resp = client.post(
        "/api/parse-resume",
        files={"file": ("resume.pdf", b"not a real pdf file", "application/pdf")},
    )
    assert resp.status_code == 400
    detail = resp.json()["detail"]
    assert "Traceback" not in detail
    assert "Failed to read this PDF" in detail
